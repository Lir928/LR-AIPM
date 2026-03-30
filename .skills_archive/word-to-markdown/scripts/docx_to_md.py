#!/usr/bin/env python3
"""
Word (.docx) to Markdown converter
Supports: headings, bold, italic, tables, images, lists, links
Requirements: python-docx, Pillow (optional, for image handling)
"""

import os
import sys
import re
import argparse
from pathlib import Path
from typing import Optional
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE


def get_media_folder(output_path: str) -> str:
    """Get the media folder path for storing images."""
    base = os.path.splitext(output_path)[0]
    return f"{base}_media"


def save_image(image_blob: bytes, media_folder: str, image_idx: int) -> str:
    """Save image blob to media folder and return the relative path."""
    os.makedirs(media_folder, exist_ok=True)

    # Determine image format from magic bytes
    if image_blob[:2] == b'\xff\xd8':
        ext = 'jpg'
    elif image_blob[:4] == b'\x89PNG':
        ext = 'png'
    elif image_blob[:4] == b'GIF8':
        ext = 'gif'
    elif image_blob[:8] == b'RIFF' and image_blob[8:12] == b'WEBP':
        ext = 'webp'
    else:
        ext = 'png'

    filename = f"image_{image_idx:03d}.{ext}"
    filepath = os.path.join(media_folder, filename)

    with open(filepath, 'wb') as f:
        f.write(image_blob)

    return filename


def is_heading_style(style_name: str) -> tuple:
    """Check if style is a heading and return level. Returns (is_heading, level)."""
    if not style_name:
        return False, 0

    # Check for standard heading styles
    if style_name.startswith('Heading '):
        try:
            level = int(style_name.split()[-1])
            return True, max(1, min(level, 6))
        except (ValueError, IndexError):
            pass

    # Check for Chinese heading styles
    chinese_headings = {
        '标题 1': 1, '标题 2': 2, '标题 3': 3,
        '标题 4': 4, '标题 5': 5, '标题 6': 6,
        '一级标题': 1, '二级标题': 2, '三级标题': 3,
        'Heading1': 1, 'Heading2': 2, 'Heading3': 3,
    }
    if style_name in chinese_headings:
        return True, chinese_headings[style_name]

    return False, 0


def get_run_text(run) -> str:
    """Extract text from a run with inline styles."""
    text = run.text
    if not text:
        return ""

    # Apply styles in reverse order to preserve nesting
    # Bold
    if run.bold:
        text = f"**{text}**"
    # Italic
    if run.italic:
        text = f"*{text}*"
    # Underline (use HTML-style since markdown doesn't have native underline)
    if run.underline:
        text = f"<u>{text}</u>"
    # Strikethrough
    if run.font.strike:
        text = f"~~{text}~~"

    return text


def process_inline_hyperlinks(runs, text: str) -> str:
    """Process hyperlinks within paragraph text."""
    # Note: python-docx doesn't easily expose hyperlinks in runs
    # We would need to parse the XML directly for full hyperlink support
    # For now, we'll handle basic inline links differently
    return text


def has_inline_image(run) -> bool:
    """Check if a run contains an inline image (drawing)."""
    if not run._element:
        return False
    # Check for drawing element
    drawing = run._element.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing')
    return drawing is not None


def get_paragraph_text(para: Paragraph, media_folder: str = None, image_map: dict = None) -> tuple:
    """
    Extract text from paragraph with full style preservation.
    Returns: (text_content, list of image references found in this paragraph)
    """
    images_in_para = []

    if not para.text.strip() and not any(has_inline_image(run) for run in para.runs):
        return "", images_in_para

    # Get style name
    style_name = para.style.name if para.style else ""

    # Check for heading
    is_heading, level = is_heading_style(style_name)
    if is_heading:
        return f"{'#' * level} {para.text.strip()}", images_in_para

    # Process runs with inline styles
    runs_text = []
    for run in para.runs:
        # Check if this run has an inline image
        if has_inline_image(run) and image_map:
            # Find the image reference in our map
            # We need to match by the drawing element
            drawing = run._element.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing')
            if drawing is not None:
                # Try to find a matching image in our map
                # Use a placeholder approach - we'll replace these after
                img_placeholder = f"__IMG_PLACEHOLDER_{id(drawing)}__"
                runs_text.append(img_placeholder)
                images_in_para.append({
                    'placeholder': img_placeholder,
                    'drawing_id': id(drawing)
                })
        else:
            runs_text.append(get_run_text(run))

    result = "".join(runs_text)

    return result, images_in_para


def is_list_paragraph(para: Paragraph) -> tuple:
    """Check if paragraph is part of a list and return (is_list, list_type, level)."""
    style_name = para.style.name if para.style else ""

    # Check for list/bullet styles
    if 'List' in style_name or 'Bullet' in style_name:
        # Determine if it's numbered or bulleted
        if 'Number' in style_name or 'Num' in style_name:
            return True, 'ordered', 0
        return True, 'unordered', 0

    # Check paragraph numbering
    if para._p is not None:
        pPr = para._p.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr')
        if pPr is not None:
            numPr = pPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr')
            if numPr is not None:
                # Check for list type
                numId = numPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numId')
                ilvl = numPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ilvl')

                level = 0
                if ilvl is not None:
                    val = ilvl.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
                    if val:
                        level = int(val)

                return True, 'numbered', level

    return False, None, 0


def convert_table(table: Table) -> str:
    """Convert Word table to Markdown table."""
    if not table.rows:
        return ""

    lines = []

    # Process each row
    for row_idx, row in enumerate(table.rows):
        cells = []
        for cell in row.cells:
            # Get cell text, normalize whitespace, escape pipes
            text = cell.text.strip()
            text = text.replace('\n', ' ').replace('\r', ' ')
            text = text.replace('|', '\\|')
            cells.append(text)

        # Create row
        lines.append("| " + " | ".join(cells) + " |")

        # Add separator after header row
        if row_idx == 0:
            lines.append("| " + " | ".join(['---'] * len(cells)) + " |")

    return "\n".join(lines)


def extract_drawing_blip_id(drawing_elem) -> str:
    """Extract the blip embed ID from a drawing element."""
    if drawing_elem is None:
        return None
    # Look for blip element which contains the embed reference
    blip = drawing_elem.find('.//{http://schemas.openxmlformats.org/drawingml/2006/main}blip')
    if blip is not None:
        embed = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
        return embed
    return None


def build_image_relationship_map(doc: Document) -> dict:
    """
    Build a map of relationship IDs to image data.
    Returns: {rId: {'data': bytes, 'filename': str}}
    """
    rel_map = {}
    image_idx = 0

    for rel_id, rel in doc.part.rels.items():
        if "image" in rel.reltype:
            try:
                image_data = rel.target_part.blob
                image_idx += 1

                # Determine image format from magic bytes
                if image_data[:2] == b'\xff\xd8':
                    ext = 'jpg'
                elif image_data[:4] == b'\x89PNG':
                    ext = 'png'
                elif image_data[:4] == b'GIF8':
                    ext = 'gif'
                elif image_data[:8] == b'RIFF' and image_data[8:12] == b'WEBP':
                    ext = 'webp'
                else:
                    ext = 'png'

                filename = f"image_{image_idx:03d}.{ext}"
                rel_map[rel_id] = {
                    'data': image_data,
                    'filename': filename,
                    'index': image_idx
                }
            except Exception as e:
                print(f"Warning: Could not extract image: {e}")

    return rel_map


def find_images_in_paragraph(para: Paragraph) -> list:
    """
    Find all images in a paragraph and their position in the text.
    Returns: list of {'rId': str, 'position': int (character position)}
    """
    images = []
    char_pos = 0

    for run in para.runs:
        # Check if this run has a drawing element
        drawing = run._element.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing')
        if drawing is not None:
            rId = extract_drawing_blip_id(drawing)
            if rId:
                images.append({
                    'rId': rId,
                    'position': char_pos
                })
        # Add run text length to position (even if it's just the placeholder)
        char_pos += len(run.text) if run.text else 0

    return images


def process_document(doc: Document, media_folder: str) -> tuple:
    """
    Process entire document and return Markdown content.
    Returns: (markdown_content, list of image info)
    """
    markdown_lines = []

    # Track list state
    in_list = False
    list_type = None

    # Build image relationship map first
    image_rel_map = build_image_relationship_map(doc)
    image_list = list(image_rel_map.values())

    # Process all block-level elements
    for element in doc.element.body:
        tag = element.tag.split('}')[-1] if '}' in element.tag else element.tag

        if tag == 'p':
            # Find the corresponding paragraph object
            para = None
            for p in doc.paragraphs:
                if p._element == element:
                    para = p
                    break

            if not para:
                continue

            # Check for images in this paragraph
            images_in_para = find_images_in_paragraph(para)

            is_list, list_kind, level = is_list_paragraph(para)

            if is_list:
                # Handle list items
                indent = "  " * level

                # Get raw text for list (without heading formatting)
                raw_text = para.text.strip()

                # If there are images, insert them after the list item
                if images_in_para:
                    list_content = raw_text if raw_text else ""
                    for img_info in images_in_para:
                        rId = img_info['rId']
                        if rId in image_rel_map:
                            img_filename = image_rel_map[rId]['filename']
                            rel_path = f"{os.path.basename(media_folder)}/{img_filename}"
                            list_content += f"\n\n{indent}![Image]({rel_path})"

                    if list_kind == 'unordered':
                        item_text = f"{indent}- {list_content.strip()}"
                    else:
                        item_text = f"{indent}1. {list_content.strip()}"
                else:
                    for run in para.runs:
                        raw_text = raw_text.replace(run.text, get_run_text(run))

                    if list_kind == 'unordered':
                        item_text = f"{indent}- {raw_text}"
                    else:
                        item_text = f"{indent}1. {raw_text}"

                markdown_lines.append(item_text)
                in_list = True

            else:
                # Check for heading
                style_name = para.style.name if para.style else ""
                is_heading, level = is_heading_style(style_name)

                if is_heading:
                    # Headings don't contain images typically
                    text = f"{'#' * level} {para.text.strip()}"
                    if in_list:
                        markdown_lines.append("")
                        in_list = False
                    markdown_lines.append(text)
                    markdown_lines.append("")

                elif para.text.strip() or images_in_para:
                    # Add spacing if transitioning from list
                    if in_list:
                        markdown_lines.append("")
                        in_list = False

                    # Process text with inline styles
                    runs_text = []
                    for run in para.runs:
                        drawing = run._element.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing')
                        if drawing is not None:
                            # Skip image placeholders in text
                            continue
                        runs_text.append(get_run_text(run))

                    text = "".join(runs_text).strip()

                    if text:
                        markdown_lines.append(text)

                    # Add images after the paragraph text
                    for img_info in images_in_para:
                        rId = img_info['rId']
                        if rId in image_rel_map:
                            img_filename = image_rel_map[rId]['filename']
                            rel_path = f"{os.path.basename(media_folder)}/{img_filename}"
                            markdown_lines.append(f"\n![Image]({rel_path})")

                    markdown_lines.append("")

        elif tag == 'tbl':
            # Find the corresponding table object
            table = None
            for t in doc.tables:
                if t._element == element:
                    table = t
                    break

            if table:
                if in_list:
                    markdown_lines.append("")
                    in_list = False

                table_md = convert_table(table)
                if table_md:
                    markdown_lines.append(table_md)
                    markdown_lines.append("")

    # Build markdown content
    markdown_content = "\n".join(markdown_lines)

    # Clean up extra blank lines
    markdown_content = re.sub(r'\n{3,}', '\n\n', markdown_content)

    return markdown_content, image_list


def docx_to_md(input_path: str, output_path: Optional[str] = None) -> tuple:
    """
    Convert Word document to Markdown.

    Returns: (output_path, media_folder_path)
    """
    # Validate input file
    input_path = os.path.abspath(input_path)
    if not os.path.exists(input_path):
        raise FileNotFoundError(f"Input file not found: {input_path}")

    if not input_path.lower().endswith('.docx'):
        raise ValueError("Input file must be a .docx file")

    # Determine output path
    if output_path is None:
        output_path = os.path.splitext(input_path)[0] + '.md'
    else:
        output_path = os.path.abspath(output_path)

    # Determine media folder
    media_folder = get_media_folder(output_path)

    # Load document
    doc = Document(input_path)

    # Convert
    markdown_content, image_list = process_document(doc, media_folder)

    # Save images
    if image_list:
        os.makedirs(media_folder, exist_ok=True)
        for img_info in image_list:
            filepath = os.path.join(media_folder, img_info['filename'])
            with open(filepath, 'wb') as f:
                f.write(img_info['data'])

    # Write output file
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(markdown_content)

    return output_path, media_folder


def main():
    parser = argparse.ArgumentParser(
        description='Convert Word (.docx) to Markdown (.md)',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog='''
Examples:
  python docx_to_md.py input.docx
  python docx_to_md.py input.docx output.md
'''
    )
    parser.add_argument('input', help='Input .docx file')
    parser.add_argument('output', nargs='?', help='Output .md file (default: same name as input)')

    args = parser.parse_args()

    try:
        output_path, media_folder = docx_to_md(args.input, args.output)

        print(f"Conversion complete!")
        print(f"  Markdown: {output_path}")

        # Check if media folder exists and has images
        if os.path.exists(media_folder):
            image_count = len([f for f in os.listdir(media_folder) if os.path.isfile(os.path.join(media_folder, f))])
            if image_count > 0:
                print(f"  Images: {media_folder}/ ({image_count} images)")

    except FileNotFoundError as e:
        print(f"Error: {e}", file=sys.stderr)
        sys.exit(1)
    except Exception as e:
        print(f"Error during conversion: {e}", file=sys.stderr)
        sys.exit(1)


if __name__ == '__main__':
    main()
