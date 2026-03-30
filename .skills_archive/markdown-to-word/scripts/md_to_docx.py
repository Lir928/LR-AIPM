#!/usr/bin/env python3
"""
Markdown to Word Document Converter (Pure Python)

Converts Markdown files to Word (.docx) format using python-docx.
No external dependencies required beyond python-docx.

Usage:
    python md_to_docx.py <input.md> [output.docx]
    python md_to_docx.py <input.md> --output-dir <directory>

Examples:
    python md_to_docx.py report.md
    python md_to_docx.py report.md output.docx
    python md_to_docx.py report.md --output-dir ./documents
"""

import argparse
import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("Error: python-docx is required. Install with: pip install python-docx")
    sys.exit(1)


class MarkdownToWord:
    """Convert Markdown to Word document."""

    def __init__(self, input_path: str, output_path: str = None):
        self.input_path = Path(input_path)
        self.output_path = Path(output_path) if output_path else self.input_path.with_suffix('.docx')
        self.doc = Document()
        self.base_dir = self.input_path.parent

        # Track code block state
        self.in_code_block = False
        self.code_block_content = []
        self.code_block_lang = ""

    def convert(self) -> str:
        """Convert Markdown file to Word document."""
        if not self.input_path.exists():
            raise FileNotFoundError(f"Input file not found: {self.input_path}")

        with open(self.input_path, 'r', encoding='utf-8') as f:
            content = f.read()

        self._parse_markdown(content)
        self._save()
        return str(self.output_path)

    def _parse_markdown(self, content: str):
        """Parse Markdown content and convert to Word."""
        lines = content.split('\n')
        i = 0

        while i < len(lines):
            line = lines[i]

            # Code block
            if line.strip().startswith('```'):
                if not self.in_code_block:
                    self.in_code_block = True
                    self.code_block_lang = line.strip()[3:].strip()
                    self.code_block_content = []
                else:
                    self.in_code_block = False
                    self._add_code_block(self.code_block_content, self.code_block_lang)
                i += 1
                continue

            if self.in_code_block:
                self.code_block_content.append(line)
                i += 1
                continue

            # Empty line
            if not line.strip():
                i += 1
                continue

            # Heading
            if line.startswith('#'):
                self._add_heading(line)
                i += 1
                continue

            # Table
            if '|' in line and i + 1 < len(lines) and '|' in lines[i + 1]:
                table_lines = [line]
                j = i + 1
                while j < len(lines) and '|' in lines[j]:
                    table_lines.append(lines[j])
                    j += 1
                self._add_table(table_lines)
                i = j
                continue

            # Unordered list
            if line.strip().startswith(('- ', '* ', '+ ')):
                self._add_list(line, ordered=False)
                i += 1
                continue

            # Ordered list
            if re.match(r'^\s*\d+\.\s', line):
                self._add_list(line, ordered=True)
                i += 1
                continue

            # Image
            img_match = re.match(r'!\[([^\]]*)\]\(([^)]+)\)', line.strip())
            if img_match:
                self._add_image(img_match.group(2), img_match.group(1))
                i += 1
                continue

            # Horizontal rule
            if re.match(r'^[-*_]{3,}\s*$', line.strip()):
                self._add_horizontal_rule()
                i += 1
                continue

            # Paragraph
            self._add_paragraph(line)
            i += 1

    def _add_heading(self, line: str):
        """Add heading to document."""
        match = re.match(r'^(#{1,6})\s+(.+)$', line)
        if not match:
            return

        level = len(match.group(1))
        text = match.group(2).strip()
        self.doc.add_heading(text, level=level)

    def _add_paragraph(self, line: str):
        """Add paragraph with inline formatting."""
        p = self.doc.add_paragraph()
        self._add_formatted_text(p, line)

    def _add_formatted_text(self, paragraph, text: str):
        """Add text with inline formatting (bold, italic, code, links)."""
        # Pattern for inline formatting - order matters: *** (bold+italic) first, then ** (bold), then * (italic)
        # Use [\s\S]? instead of .? to properly match content, handle edge cases where ** appears without closing
        patterns = [
            (r'\*\*\*(.+?)\*\*\*', {'bold': True, 'italic': True}),
            (r'\*\*(.+?)\*\*', {'bold': True}),
            (r'\*(.+?)\*', {'italic': True}),
            (r'~~(.+?)~~', {'strikethrough': True}),
            (r'`([^`]+)`', {'code': True}),
            (r'\[([^\]]+)\]\(([^)]+)\)', {'link': True}),
        ]

        # Build combined pattern
        combined = '|'.join(f'({p[0]})' for p in patterns)

        last_end = 0
        for match in re.finditer(combined, text):
            # Add text before match
            if match.start() > last_end:
                run = paragraph.add_run(text[last_end:match.start()])

            # Find which group matched
            for idx, (pattern, fmt) in enumerate(patterns):
                outer_group = idx + 1  # Outer group captures the full match including delimiters
                inner_group = idx + 2   # Inner group captures just the content
                try:
                    # Check if this pattern's outer group matched
                    if match.group(outer_group):
                        if fmt.get('link'):
                            # Link - get text from outer group, URL from inner group
                            link_text = match.group(outer_group)
                            # Extract text from [text](url) format
                            link_match = re.match(r'\[([^\]]+)\]\(([^)]+)\)', link_text)
                            if link_match:
                                link_text = link_match.group(1)
                                link_url = link_match.group(2)
                            else:
                                link_url = link_text
                            run = paragraph.add_run(link_text)
                            run.font.color.rgb = RGBColor(0, 0, 255)
                            run.font.underline = True
                        elif fmt.get('code'):
                            # Code - use outer group (includes backticks)
                            code_content = match.group(outer_group)
                            run = paragraph.add_run(code_content)
                            run.font.name = 'Consolas'
                            run.font.size = Pt(9)
                        else:
                            # Bold, italic, strikethrough - use inner group (content only)
                            content = match.group(inner_group)
                            if content:
                                run = paragraph.add_run(content)
                                if fmt.get('bold'):
                                    run.bold = True
                                if fmt.get('italic'):
                                    run.italic = True
                                if fmt.get('strikethrough'):
                                    run.font.strike = True
                        break
                except (IndexError, AttributeError):
                    continue

            last_end = match.end()

        # Add remaining text
        if last_end < len(text):
            paragraph.add_run(text[last_end:])

    def _add_list(self, line: str, ordered: bool = False):
        """Add list item to document."""
        # Parse indent level and content
        match = re.match(r'^(\s*)([-*+]+|\d+\.)\s+(.+)$', line)
        if not match:
            return

        indent = len(match.group(1)) // 2
        content = match.group(3)

        # Add paragraph for list item
        p = self.doc.add_paragraph(style='List Bullet' if not ordered else 'List Number')

        # Handle nested lists with proper indent
        if indent > 0:
            p.paragraph_format.left_indent = Inches(0.5 * indent)

        self._add_formatted_text(p, content)

    def _add_table(self, lines: list):
        """Add table to document."""
        if len(lines) < 2:
            return

        # Parse table rows
        rows = []
        for line in lines:
            # Skip separator lines (handle various formats like |--|, |---|, |:-:|, etc.)
            if re.match(r'^\s*\|?[\s\-:|]+\|?\s*$', line):
                continue
            cells = [cell.strip() for cell in line.split('|') if cell.strip() or line.count('|') > 1]
            # Re-parse to handle empty cells properly
            if '|' in line:
                parts = line.split('|')
                # Remove first and last empty parts (from leading/trailing |)
                if parts and not parts[0].strip():
                    parts = parts[1:]
                if parts and not parts[-1].strip():
                    parts = parts[:-1]
                cells = [p.strip() for p in parts]
            if cells:
                rows.append(cells)

        if not rows:
            return

        # Create table
        num_cols = max(len(row) for row in rows)
        table = self.doc.add_table(rows=len(rows), cols=num_cols)
        table.style = 'Table Grid'

        for i, row_data in enumerate(rows):
            for j, cell_text in enumerate(row_data):
                if j < num_cols:
                    cell = table.rows[i].cells[j]
                    # Clear default paragraph and add formatted text
                    cell.paragraphs[0].clear()
                    self._add_formatted_text(cell.paragraphs[0], cell_text)
                    # Bold header row
                    if i == 0:
                        for run in cell.paragraphs[0].runs:
                            run.bold = True

    def _add_code_block(self, lines: list, lang: str = ""):
        """Add code block to document."""
        if not lines:
            return

        # Add language hint if present
        if lang:
            p = self.doc.add_paragraph()
            run = p.add_run(f"[{lang}]")
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(128, 128, 128)

        # Add code content
        code_text = '\n'.join(lines)
        p = self.doc.add_paragraph()
        run = p.add_run(code_text)
        run.font.name = 'Consolas'
        run.font.size = Pt(9)

        # Add background color (via shading)
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), 'F5F5F5')
        p._p.get_or_add_pPr().append(shading)

    def _add_image(self, img_path: str, alt_text: str = ""):
        """Add image to document."""
        full_path = self.base_dir / img_path
        if full_path.exists():
            try:
                p = self.doc.add_paragraph()
                run = p.add_run()
                run.add_picture(str(full_path), width=Inches(5))
                if alt_text:
                    cap = self.doc.add_paragraph(alt_text)
                    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in cap.runs:
                        run.italic = True
                        run.font.size = Pt(9)
            except Exception as e:
                self.doc.add_paragraph(f"[Image: {img_path}]")
        else:
            self.doc.add_paragraph(f"[Image not found: {img_path}]")

    def _add_horizontal_rule(self):
        """Add horizontal rule to document."""
        p = self.doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        # Add a border as horizontal rule
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), 'auto')
        pBdr.append(bottom)
        p._p.get_or_add_pPr().append(pBdr)

    def _save(self):
        """Save the document."""
        self.output_path.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(str(self.output_path))


def main():
    parser = argparse.ArgumentParser(
        description="Convert Markdown files to Word documents",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  %(prog)s report.md                    # Convert to report.docx
  %(prog)s report.md output.docx        # Convert to specified output
  %(prog)s report.md --output-dir docs  # Save to docs/ directory
        """
    )
    parser.add_argument("input", help="Input Markdown file")
    parser.add_argument("output", nargs="?", help="Output Word file")
    parser.add_argument("--output-dir", "-d", help="Output directory")

    args = parser.parse_args()

    try:
        input_path = Path(args.input)

        if args.output:
            output_path = args.output
        elif args.output_dir:
            output_path = Path(args.output_dir) / f"{input_path.stem}.docx"
        else:
            output_path = None

        converter = MarkdownToWord(str(input_path), str(output_path) if output_path else None)
        result = converter.convert()
        print(f"Created: {result}")

    except FileNotFoundError as e:
        print(f"Error: {e}", file=sys.stderr)
        sys.exit(1)
    except Exception as e:
        print(f"Error: {e}", file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    main()